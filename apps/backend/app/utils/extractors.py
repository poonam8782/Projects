"""
Text extraction utilities for Neura.

This module provides text extraction functions for different file types:
- PDF files using PyMuPDF (fitz)
- DOCX files using python-docx
- Image files (PNG, JPG) using pytesseract OCR
"""

import logging
from io import BytesIO
from typing import Optional, Callable

import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        file_bytes: Raw PDF file bytes
        
    Returns:
        Extracted text as a string, or empty string on failure
    """
    try:
        # Open PDF from bytes
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            text_parts = []
            for page in doc:
                # Extract text in plain text mode
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(page_text)
            
            # Join pages with double newline for readability
            return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    Extracts both paragraph text and table content.
    
    Args:
        file_bytes: Raw DOCX file bytes
        
    Returns:
        Extracted text as a string, or empty string on failure
    """
    try:
        # Open document from bytes
        doc = Document(BytesIO(file_bytes))
        
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cell_texts = []
                for cell in row.cells:
                    # Join cell paragraphs with newline
                    cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                    if cell_text:
                        cell_texts.append(cell_text)
                if cell_texts:
                    # Join cells with tab separator
                    table_rows.append("\t".join(cell_texts))
            if table_rows:
                # Join rows with newline
                text_parts.append("\n".join(table_rows))
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        return ""


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from an image file using pytesseract OCR.
    
    Args:
        file_bytes: Raw image file bytes (PNG, JPG, etc.)
        
    Returns:
        Extracted text as a string, or error message on failure
    """
    try:
        # Open image from bytes
        img = Image.open(BytesIO(file_bytes))
        
        # Convert to RGB if needed (some PNGs have alpha channel)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(img, lang="eng")
        return text.strip()
    except pytesseract.TesseractNotFoundError:
        logger.error("Tesseract OCR not found on system")
        return ""
    except Exception as e:
        logger.error(f"Image text extraction failed: {e}")
        return ""


def get_extractor_for_mime_type(mime_type: str) -> Optional[Callable[[bytes], str]]:
    """
    Factory function to get the appropriate text extractor for a given MIME type.
    
    Args:
        mime_type: MIME type of the file
        
    Returns:
        Extractor function that takes bytes and returns string, or None if no extractor available
    """
    extractors = {
        "application/pdf": extract_text_from_pdf,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": extract_text_from_docx,
        "text/plain": lambda file_bytes: file_bytes.decode("utf-8", errors="ignore"),
        "image/png": extract_text_from_image,
        "image/jpeg": extract_text_from_image,
    }
    
    return extractors.get(mime_type)
